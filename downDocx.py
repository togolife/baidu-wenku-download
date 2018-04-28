#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import urllib2
import json
import sys
reload(sys)
sys.setdefaultencoding('utf-8')

# word操作库python-docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
# python-docx 格式参考：https://zhuanlan.zhihu.com/p/23708800?utm_source=tuicool&utm_medium=referral

import config
import log

logger = log.Log(config.log_dir, config.log_name)

class DownDocx(object):
	def __init__(self, fileDir, url, info):
		self.fileDir = './' + fileDir + '/'
		self.URL = url
		self.WkInfo = info
		if not os.path.exists(self.fileDir):
			os.mkdir(self.fileDir)

	# 获取每页URL
	@staticmethod
	def geturl(urls, index):
		for url in urls:
			if url['pageIndex'] == index:
				return url['pageLoadUrl']
		return ''

	def down(self):
		reqHeader = config.reqHeaderBDWK
		reqHeader['Referer'] = self.URL
		i = 1
		docFileName = self.fileDir + self.WkInfo['title'] + '.' + self.WkInfo['docType']
		document = Document() # 创建doc(x)文档
		while i <= self.WkInfo['totalPageNum']:
			jsonUrl = DownDocx.geturl(self.WkInfo['htmlUrls']['json'], i)
			if len(jsonUrl) == 0:
				logger.error('下载文档失败，查找URL失败！')
				return False
			jsonUrl = jsonUrl.replace('\\', '')
			jsonUrl = jsonUrl.replace(' ', '%20')
			req = urllib2.Request(jsonUrl, headers=reqHeader)
			res = urllib2.urlopen(req)
			res = res.read()
			jsonRet = res[res.find('(')+1 : res.rfind(')')]
			logger.info('打印一下，获取json数据内容为 ' + jsonRet)
			jsonRet = json.loads(jsonRet)
			# 再处理获取的页面内容
			first = 0
			for item in jsonRet['body']:
				if item['t'] != 'word':
					continue
				if first == 0 or (item['ps'] and item['ps']['_enter'] == 1):
					first = 1
					pg = document.add_paragraph()
				if item['ps'] and item['ps']['_enter'] == 1:
					continue
				run = pg.add_run(item['c'])
				# 添加格式；分析不出来，就统一宋体、五号
				run.font.name = u'宋体'
				run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
				run.font.size = Pt(10.5)
			# 下一页
			if i < self.WkInfo['totalPageNum']:
				document.add_page_break()
			i += 1
		document.save(docFileName)
		return True